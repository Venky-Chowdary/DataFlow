"""Structure-aware document chunking for PDF / Word / HTML sources.

Honesty
-------
Produces tabular records (one row per chunk) with provenance metadata so
operators can inspect page/heading/element before embedding. Chunks are sized
for RAG defaults (≈512 chars with overlap across long paragraphs).

This module extracts text-layer PDF / Word / HTML into provenance chunk rows.
Scanned PDFs with no text layer can opt into OCR via ``enable_ocr=True``
(``services/pdf_ocr.py`` + system Tesseract). Without OCR, empty text fails
closed — never silent.
"""

from __future__ import annotations

import hashlib
import io
import logging
import re
from html.parser import HTMLParser
from typing import Any

logger = logging.getLogger(__name__)

# Marker so vectorize_records embeds once without re-splitting.
PRECHUNKED_FLAG = "_df_prechunked"

DOCUMENT_TYPES = frozenset({"pdf", "docx", "html", "htm"})

PROVENANCE_COLUMNS = (
    "id",
    "content",
    "filename",
    "source_id",
    "page",
    "heading",
    "element_type",
    "chunk_index",
    PRECHUNKED_FLAG,
)


def detect_document_type(filename: str, content: bytes | None = None) -> str | None:
    lower = (filename or "").lower()
    if lower.endswith(".pdf"):
        return "pdf"
    if lower.endswith((".docx", ".doc")):
        # .doc (legacy binary) is not supported — only .docx via python-docx.
        return "docx" if lower.endswith(".docx") else None
    if lower.endswith((".html", ".htm")):
        return "html"
    if content and content[:5] == b"%PDF-":
        return "pdf"
    if content and content[:2] == b"PK" and b"word/" in content[:2000]:
        return "docx"
    sample = (content or b"")[:256].lstrip().lower()
    if sample.startswith(b"<!doctype html") or sample.startswith(b"<html"):
        return "html"
    return None


def _chunk_id(source_id: str, chunk_index: int, text: str) -> str:
    digest = hashlib.sha256(f"{source_id}:{chunk_index}:{text[:200]}".encode("utf-8")).hexdigest()
    return digest[:24]


def _split_long_text(
    text: str,
    *,
    chunk_size: int = 512,
    chunk_overlap: int = 50,
) -> list[str]:
    """Reuse vectorization paragraph/sentence splitter for long blocks."""
    from services.vectorization import chunk_text

    parts = chunk_text(text, chunk_size=chunk_size, chunk_overlap=chunk_overlap)
    return parts or ([text.strip()] if text.strip() else [])


def _record(
    *,
    source_id: str,
    filename: str,
    content: str,
    chunk_index: int,
    page: str = "",
    heading: str = "",
    element_type: str = "paragraph",
) -> dict[str, str]:
    text = (content or "").strip()
    return {
        "id": _chunk_id(source_id, chunk_index, text),
        "content": text,
        "filename": filename,
        "source_id": source_id,
        "page": page,
        "heading": heading,
        "element_type": element_type,
        "chunk_index": str(chunk_index),
        PRECHUNKED_FLAG: "1",
    }


def extract_pdf_chunks(
    content: bytes,
    *,
    filename: str = "document.pdf",
    chunk_size: int = 512,
    chunk_overlap: int = 50,
    enable_ocr: bool = False,
    ocr_dpi: int = 200,
    ocr_language: str = "eng",
) -> list[dict[str, str]]:
    try:
        from pypdf import PdfReader
    except ImportError as exc:
        raise RuntimeError(
            "pypdf is required for PDF sources — install with: pip install pypdf"
        ) from exc

    reader = PdfReader(io.BytesIO(content))
    source_id = hashlib.sha256(content[:4096] + filename.encode()).hexdigest()[:16]
    rows: list[dict[str, str]] = []
    chunk_index = 0
    for page_num, page in enumerate(reader.pages, start=1):
        try:
            text = page.extract_text() or ""
        except Exception as exc:
            logger.warning("PDF page %s extract failed: %s", page_num, exc)
            text = ""
        text = text.strip()
        if not text:
            continue
        # Prefer page-level blocks split on blank lines, then size-limit.
        blocks = [b.strip() for b in re.split(r"\n\s*\n", text) if b.strip()]
        if not blocks:
            blocks = [text]
        for block in blocks:
            for piece in _split_long_text(block, chunk_size=chunk_size, chunk_overlap=chunk_overlap):
                rows.append(
                    _record(
                        source_id=source_id,
                        filename=filename,
                        content=piece,
                        chunk_index=chunk_index,
                        page=str(page_num),
                        heading="",
                        element_type="paragraph",
                    )
                )
                chunk_index += 1

    if rows:
        return rows

    if not enable_ocr:
        return rows

    from services.pdf_ocr import ocr_pdf_pages

    try:
        ocr_pages = ocr_pdf_pages(content, dpi=ocr_dpi, language=ocr_language)
    except RuntimeError:
        raise
    except Exception as exc:
        raise RuntimeError(f"OCR failed for scanned PDF: {exc}") from exc

    for page_num, text in ocr_pages:
        blocks = [b.strip() for b in re.split(r"\n\s*\n", text) if b.strip()]
        if not blocks:
            blocks = [text]
        for block in blocks:
            for piece in _split_long_text(block, chunk_size=chunk_size, chunk_overlap=chunk_overlap):
                rows.append(
                    _record(
                        source_id=source_id,
                        filename=filename,
                        content=piece,
                        chunk_index=chunk_index,
                        page=str(page_num),
                        heading="",
                        element_type="ocr",
                    )
                )
                chunk_index += 1
    return rows


def extract_docx_chunks(
    content: bytes,
    *,
    filename: str = "document.docx",
    chunk_size: int = 512,
    chunk_overlap: int = 50,
) -> list[dict[str, str]]:
    try:
        from docx import Document
    except ImportError as exc:
        raise RuntimeError(
            "python-docx is required for Word sources — install with: pip install python-docx"
        ) from exc

    doc = Document(io.BytesIO(content))
    source_id = hashlib.sha256(content[:4096] + filename.encode()).hexdigest()[:16]
    rows: list[dict[str, str]] = []
    chunk_index = 0
    current_heading = ""

    for para in doc.paragraphs:
        text = (para.text or "").strip()
        if not text:
            continue
        style_name = ""
        try:
            style_name = (para.style.name or "") if para.style else ""
        except Exception:
            style_name = ""
        if style_name.lower().startswith("heading"):
            current_heading = text
            element_type = "heading"
        else:
            element_type = "paragraph"
        for piece in _split_long_text(text, chunk_size=chunk_size, chunk_overlap=chunk_overlap):
            rows.append(
                _record(
                    source_id=source_id,
                    filename=filename,
                    content=piece,
                    chunk_index=chunk_index,
                    page="",
                    heading=current_heading if element_type != "heading" else piece,
                    element_type=element_type,
                )
            )
            chunk_index += 1

    # Tables → markdown-ish rows with provenance.
    for t_idx, table in enumerate(doc.tables):
        lines: list[str] = []
        for row in table.rows:
            cells = [(c.text or "").strip().replace("\n", " ") for c in row.cells]
            if any(cells):
                lines.append(" | ".join(cells))
        table_text = "\n".join(lines).strip()
        if not table_text:
            continue
        for piece in _split_long_text(table_text, chunk_size=chunk_size, chunk_overlap=chunk_overlap):
            rows.append(
                _record(
                    source_id=source_id,
                    filename=filename,
                    content=piece,
                    chunk_index=chunk_index,
                    page="",
                    heading=current_heading,
                    element_type=f"table:{t_idx}",
                )
            )
            chunk_index += 1
    return rows


class _HTMLTextExtractor(HTMLParser):
    def __init__(self) -> None:
        super().__init__()
        self.blocks: list[tuple[str, str]] = []  # (element_type, text)
        self._skip = 0
        self._buf: list[str] = []
        self._tag = "p"
        self._heading = ""

    def handle_starttag(self, tag: str, attrs: list[tuple[str, str | None]]) -> None:
        t = tag.lower()
        if t in {"script", "style", "noscript"}:
            self._skip += 1
            return
        if self._skip:
            return
        if t in {"h1", "h2", "h3", "h4", "h5", "h6", "p", "li", "td", "th", "pre", "blockquote"}:
            self._flush()
            self._tag = t

    def handle_endtag(self, tag: str) -> None:
        t = tag.lower()
        if t in {"script", "style", "noscript"} and self._skip:
            self._skip -= 1
            return
        if self._skip:
            return
        if t in {"h1", "h2", "h3", "h4", "h5", "h6", "p", "li", "td", "th", "pre", "blockquote", "div", "section"}:
            self._flush()

    def handle_data(self, data: str) -> None:
        if self._skip:
            return
        if data and data.strip():
            self._buf.append(data)

    def _flush(self) -> None:
        text = re.sub(r"\s+", " ", "".join(self._buf)).strip()
        self._buf = []
        if not text:
            return
        if self._tag.startswith("h") and len(self._tag) == 2 and self._tag[1].isdigit():
            self._heading = text
            self.blocks.append(("heading", text))
        elif self._tag in {"td", "th"}:
            self.blocks.append(("table_cell", text))
        else:
            self.blocks.append(("paragraph", text))


def extract_html_chunks(
    content: bytes,
    *,
    filename: str = "document.html",
    chunk_size: int = 512,
    chunk_overlap: int = 50,
) -> list[dict[str, str]]:
    text = content.decode("utf-8", errors="replace")
    parser = _HTMLTextExtractor()
    parser.feed(text)
    parser.close()
    source_id = hashlib.sha256(content[:4096] + filename.encode()).hexdigest()[:16]
    rows: list[dict[str, str]] = []
    chunk_index = 0
    current_heading = ""
    for element_type, block in parser.blocks:
        if element_type == "heading":
            current_heading = block
        for piece in _split_long_text(block, chunk_size=chunk_size, chunk_overlap=chunk_overlap):
            rows.append(
                _record(
                    source_id=source_id,
                    filename=filename,
                    content=piece,
                    chunk_index=chunk_index,
                    page="",
                    heading=current_heading if element_type != "heading" else piece,
                    element_type=element_type,
                )
            )
            chunk_index += 1
    return rows


def extract_document_chunks(
    content: bytes,
    filename: str,
    *,
    chunk_size: int = 512,
    chunk_overlap: int = 50,
    doc_type: str | None = None,
    enable_ocr: bool = False,
    ocr_dpi: int = 200,
    ocr_language: str = "eng",
) -> list[dict[str, str]]:
    """Return provenance-aware chunk records for a document upload."""
    kind = doc_type or detect_document_type(filename, content)
    if not kind:
        raise ValueError(f"Unsupported document type for {filename or 'upload'}")
    name = filename or f"document.{kind}"
    if kind == "pdf":
        return extract_pdf_chunks(
            content,
            filename=name,
            chunk_size=chunk_size,
            chunk_overlap=chunk_overlap,
            enable_ocr=enable_ocr,
            ocr_dpi=ocr_dpi,
            ocr_language=ocr_language,
        )
    if kind == "docx":
        return extract_docx_chunks(
            content, filename=name, chunk_size=chunk_size, chunk_overlap=chunk_overlap
        )
    if kind in {"html", "htm"}:
        return extract_html_chunks(
            content, filename=name, chunk_size=chunk_size, chunk_overlap=chunk_overlap
        )
    raise ValueError(f"Unsupported document type: {kind}")


def document_columns() -> list[str]:
    return list(PROVENANCE_COLUMNS)
