"""Document chunking — real PDF/DOCX/HTML → provenance rows (no mocks)."""

from __future__ import annotations

import io
import sys
from pathlib import Path

_API_ROOT = Path(__file__).resolve().parents[1]
if str(_API_ROOT) not in sys.path:
    sys.path.insert(0, str(_API_ROOT))


def _make_pdf_bytes(text: str) -> bytes:
    """Minimal PDF with extractable Helvetica text (no reportlab required)."""
    # Escape PDF string literals.
    safe = text.replace("\\", "\\\\").replace("(", "\\(").replace(")", "\\)")[:200]
    content_stream = f"BT /F1 12 Tf 50 700 Td ({safe}) Tj ET".encode("latin-1", errors="replace")
    objects = []
    objects.append(b"1 0 obj<< /Type /Catalog /Pages 2 0 R >>endobj\n")
    objects.append(b"2 0 obj<< /Type /Pages /Kids [3 0 R] /Count 1 >>endobj\n")
    objects.append(
        b"3 0 obj<< /Type /Page /Parent 2 0 R /MediaBox [0 0 612 792] "
        b"/Contents 4 0 R /Resources << /Font << /F1 5 0 R >> >> >>endobj\n"
    )
    objects.append(
        f"4 0 obj<< /Length {len(content_stream)} >>stream\n".encode("ascii")
        + content_stream
        + b"\nendstream\nendobj\n"
    )
    objects.append(b"5 0 obj<< /Type /Font /Subtype /Type1 /BaseFont /Helvetica >>endobj\n")
    body = b"".join(objects)
    # Build xref
    header = b"%PDF-1.4\n"
    offsets = [0]
    pos = len(header)
    for obj in objects:
        offsets.append(pos)
        pos += len(obj)
    xref_pos = pos
    xref = [f"xref\n0 {len(offsets)}\n".encode("ascii")]
    xref.append(b"0000000000 65535 f \n")
    for off in offsets[1:]:
        xref.append(f"{off:010d} 00000 n \n".encode("ascii"))
    trailer = (
        f"trailer<< /Size {len(offsets)} /Root 1 0 R >>\nstartxref\n{xref_pos}\n%%EOF\n".encode(
            "ascii"
        )
    )
    return header + body + b"".join(xref) + trailer


def _make_docx_bytes(paragraphs: list[str], heading: str = "") -> bytes:
    from docx import Document

    doc = Document()
    if heading:
        doc.add_heading(heading, level=1)
    for p in paragraphs:
        doc.add_paragraph(p)
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def test_detect_document_type():
    from services.document_chunking import detect_document_type

    assert detect_document_type("a.pdf") == "pdf"
    assert detect_document_type("a.docx") == "docx"
    assert detect_document_type("a.html") == "html"
    assert detect_document_type("a.doc") is None
    assert detect_document_type("x.bin", b"%PDF-1.4") == "pdf"


def test_html_chunks_have_provenance():
    from services.document_chunking import PRECHUNKED_FLAG, extract_document_chunks
    from services.file_parser import FileParser

    html = b"""<!DOCTYPE html><html><body>
    <h1>Orders Guide</h1>
    <p>First paragraph about syncing orders safely.</p>
    <p>Second paragraph covers quarantine and replay.</p>
    </body></html>"""
    rows = extract_document_chunks(html, "guide.html")
    assert len(rows) >= 2
    assert all(r.get(PRECHUNKED_FLAG) == "1" for r in rows)
    assert any(r.get("element_type") == "heading" for r in rows)
    assert any("quarantine" in r["content"].lower() for r in rows)
    assert any(r.get("heading") for r in rows if r.get("element_type") == "paragraph")

    parsed = FileParser.parse(html, "guide.html")
    assert parsed.success is True
    assert parsed.file_type == "html"
    assert parsed.row_count >= 2
    assert "content" in parsed.columns
    assert "page" in parsed.columns


def test_docx_chunks_via_file_parser():
    from services.file_parser import FileParser

    raw = _make_docx_bytes(
        ["DataFlow moves data without silent loss.", "Quarantine keeps bad cells visible."],
        heading="Integrity",
    )
    parsed = FileParser.parse(raw, "integrity.docx")
    assert parsed.success is True
    assert parsed.file_type == "docx"
    assert parsed.row_count >= 2
    contents = " ".join(r["content"] for r in parsed.data)
    assert "silent loss" in contents
    assert any(r.get("element_type") == "heading" for r in parsed.data)


def test_prechunked_vectorize_does_not_resplit():
    """Document rows stay 1:1 through vectorize when _df_prechunked is set."""
    from services.document_chunking import PRECHUNKED_FLAG
    from services.vectorization import vectorize_records

    records = [
        {
            "id": "c0",
            "content": "Alpha chunk text for embedding.",
            "page": "1",
            "heading": "Intro",
            "chunk_index": "0",
            PRECHUNKED_FLAG: "1",
            "filename": "doc.pdf",
        },
        {
            "id": "c1",
            "content": "Beta chunk text for embedding.",
            "page": "2",
            "heading": "Intro",
            "chunk_index": "1",
            PRECHUNKED_FLAG: "1",
            "filename": "doc.pdf",
        },
    ]
    rows = vectorize_records(
        records,
        content_column="content",
        embedding_column="embedding",  # missing → embed path
        # Force precomputed empty path: use skip via flag + inject fake embedding column absent
    )
    # Without embedding_column values, it embeds — may need ST model.
    # Instead pass precomputed embeddings to stay CI-stable:
    for r in records:
        r["embedding"] = "[0.1,0.2,0.3]"
    rows = vectorize_records(records, content_column="content", embedding_column="embedding")
    assert len(rows) == 2
    assert rows[0]["chunk_index"] == 0
    assert rows[1]["chunk_index"] == 1
    assert rows[0]["metadata"].get("page") == "1"
    assert PRECHUNKED_FLAG not in rows[0]["metadata"]


def test_pdf_chunks_when_synthesizer_available():
    from services.file_parser import FileParser

    raw = _make_pdf_bytes("CDC lag and quarantine findings")
    parsed = FileParser.parse(raw, "cdc.pdf")
    assert parsed.success is True
    assert parsed.file_type == "pdf"
    assert parsed.row_count >= 1
    blob = " ".join(r["content"] for r in parsed.data)
    assert "CDC" in blob or "quarantine" in blob.lower()
    assert all(r.get("page") for r in parsed.data)
